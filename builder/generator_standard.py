from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.enum.style import WD_STYLE_TYPE
import io
import re # Import re module for URL cleaning

# Stores the calculated right margin end for consistent tab stop placement
global_margin_end = None

def _calculate_margin_end(doc, right_padding=Inches(0.12)):
    """
    Calculates the absolute right margin position for tab stop alignment (as a Length).
    Subtracts a small padding so right-aligned text won't touch the page margin.
    `right_padding` is adjustable (default 0.12") if you want more/less gap.
    """
    global global_margin_end
    if global_margin_end is None:
        section = doc.sections[0]
        # page_width and margins are Length objects; subtract a small padding so text doesn't touch edge
        right_edge = section.page_width - section.right_margin - Cm(1.27)
        global_margin_end = right_edge
    return global_margin_end

def _set_default_font_style(doc):
    """Configures the 'Normal' style's default font to Cambria 10pt and sets single line spacing."""
    styles = doc.styles
    normal_style = styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Cambria'
    normal_font.size = Pt(10)
    normal_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    normal_style.paragraph_format.space_after = Pt(2) # Reduces default spacing after paragraphs
    normal_style.paragraph_format.space_before = Pt(2) # Reduces default spacing before paragraphs

def _add_section_title(doc_obj, title_text):
    """Adds a major section title with specific formatting (12pt, bold, Cambria, and a bottom border)."""
    p = doc_obj.add_paragraph()
    p.paragraph_format.space_before = Pt(10) # Provides space above the section title
    p.paragraph_format.space_after = Pt(3) # Provides minimal space below the title text
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(title_text.upper())
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Cambria'

    # Adds a bottom border to the paragraph to visually separate sections
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6') # Border thickness
    bottom.set(qn('w:space'), '1') # Space between text and border
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)
    pPr.append(pBdr)

    p.paragraph_format.space_after = Pt(8) # Adds more space after the border to the next content block

def _get_or_create_hyperlink_style(d):
    """Ensures the 'Hyperlink' character style exists in the document.
        If missing, it creates a default version consistent with Word 2019's default theme."""
    if "Hyperlink" not in d.styles:
        # Ensures a 'Default Character Font' style exists
        if "Default Character Font" not in d.styles:
            default_char_style = d.styles.add_style("Default Character Font",
                                                     WD_STYLE_TYPE.CHARACTER,
                                                     True)
            default_char_style.element.set(qn('w:default'), "1")
            default_char_style.priority = 1
            default_char_style.hidden = True
            default_char_style.unhide_when_used = True

        # Creates the 'Hyperlink' style
        hyperlink_style = d.styles.add_style("Hyperlink",
                                              WD_STYLE_TYPE.CHARACTER,
                                              True)
        # Sets the base style and unhide it for use
        hyperlink_style.base_style = d.styles["Default Character Font"] if "Default Character Font" in d.styles else d.styles["Normal"]
        hyperlink_style.unhide_when_used = True
        hyperlink_style.font.color.rgb = RGBColor(0x05, 0x63, 0xC1) # Sets the default blue color for hyperlinks
        hyperlink_style.font.underline = True # Ensures hyperlinks are underlined

    return "Hyperlink"

def _add_hyperlink(doc_obj, paragraph, url, text, size=10):
    """
    Adds a clickable hyperlink to a given paragraph with specified text and font size.
    It handles the creation of the necessary relationship within the document.
    """
    # If the URL is not valid (empty or not HTTP/HTTPS), add plain text instead of a hyperlink
    if not url or not (url.startswith("http://") or url.startswith("https://")):
        run = paragraph.add_run(text)
        run.font.size = Pt(size)
        run.font.name = 'Cambria'
        return run

    r_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)

    # Creates the <w:hyperlink> XML element, which is the container for the clickable text.
    hyperlink_element = OxmlElement('w:hyperlink')
    hyperlink_element.set(qn('r:id'), r_id) # Links the hyperlink element to the external relationship

    new_run = paragraph.add_run(text)
    
    # Applies the 'Hyperlink' style to the run, setting its default appearance (blue, underlined).
    new_run.style = _get_or_create_hyperlink_style(doc_obj)

    # Overrides or ensures specific font properties for consistency.
    new_run.font.size = Pt(size)
    new_run.font.name = 'Cambria'
    new_run.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
    new_run.font.underline = True

    rPr = new_run._element.get_or_add_rPr()
    no_proof = OxmlElement('w:noProof')
    rPr.append(no_proof)

    parent_element = new_run._element.getparent()    
    parent_element.remove(new_run._element)
    hyperlink_element.append(new_run._element)
    parent_element.append(hyperlink_element)

    return new_run

def _clean_url_display(url):
    """Removes http(s):// and www. from a URL for cleaner display."""
    if not url:
        return ""
    # Remove protocol
    cleaned_url = re.sub(r"https?://", "", url, 1)
    # Remove www. if present at the start
    cleaned_url = re.sub(r"^www\.", "", cleaned_url, 1)
    # Remove trailing slash if present
    if cleaned_url.endswith('/'):
        cleaned_url = cleaned_url[:-1]
    return cleaned_url


def _add_left_right_paragraph(doc_obj, left_content, right_display_text, right_url=None, left_size=10, right_size=10, right_italic=False):
    """
    Adds a paragraph with content aligned to the left and optional text/hyperlink aligned to the right
    using a tab stop, with option for italic on right text.
    """
    p = doc_obj.add_paragraph()
    p.paragraph_format.space_after = Pt(2) # Minimal vertical spacing after the paragraph
    p.paragraph_format.space_before = Pt(2) # Minimal vertical spacing before the paragraph
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    # Sets a right-aligned tab stop at the calculated right margin of the content area.
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(_calculate_margin_end(doc_obj), WD_TAB_ALIGNMENT.RIGHT)

    # Processes and adds the left-aligned content.
    if isinstance(left_content, list):
        # If left_content is a list of tuples, iterate and apply specific bold/italic formatting.
        for text_part, bold_part, italic_part in left_content:
            run_left = p.add_run(text_part)
            run_left.bold = bold_part
            run_left.italic = italic_part
            run_left.font.size = Pt(left_size)
            run_left.font.name = 'Cambria'
    else: 
        # If left_content is a single string, adds it as a bold run (typical for titles).
        run_left = p.add_run(left_content)
        run_left.bold = True
        run_left.font.size = Pt(left_size)
        run_left.font.name = 'Cambria'

    # Processes and adds the right-aligned text or hyperlink.
    if right_display_text.strip():
        # Adds a tab character to push the subsequent text to the right-aligned tab stop.
        p.add_run("\t")
        if right_url:
            normalized_url = right_url.strip()
            # Consider as hyperlink if it starts with http(s) or looks like a domain/path
            if (
                normalized_url.startswith("http://")
                or normalized_url.startswith("https://")
                or "." in normalized_url
                or "/" in normalized_url
            ):
                if not (normalized_url.startswith("http://") or normalized_url.startswith("https://")):
                    normalized_url = "https://" + normalized_url
                _add_hyperlink(doc_obj, p, normalized_url, right_display_text, right_size)
            else:
                # Plain text
                run_right = p.add_run(right_display_text)
                run_right.font.size = Pt(right_size)
                run_right.font.name = 'Cambria'
                run_right.italic = right_italic
        else:
            # If no valid URL, it's plain text. Apply font and italic if requested.
            run_right = p.add_run(right_display_text)
            run_right.font.size = Pt(right_size)
            run_right.font.name = 'Cambria'
            run_right.italic = right_italic # Apply italic here


def _add_bullet_point(doc_obj, text):
    """Adds a single bullet point paragraph with Cambria 10pt font."""
    p = doc_obj.add_paragraph(text, style='List Bullet') # Uses the built-in 'List Bullet' style
    p.paragraph_format.left_indent = Inches(0.25) # Indents the bullet point
    p.paragraph_format.first_line_indent = Inches(-0.25) # Pulls the bullet symbol left
    p.paragraph_format.space_after = Pt(2) # Minimal vertical spacing after
    p.paragraph_format.space_before = Pt(2) # Minimal vertical spacing before
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    # Ensures all runs within the bullet point are Cambria 10pt
    for run in p.runs:
        run.font.name = 'Cambria'
        run.font.size = Pt(10)


def add_header_section(doc, name, title, location, email, phone, website, linkedin, github):
    """Adds the resume header including name, title, and all contact information."""
    # Paragraph for Name and Title
    name_title_para = doc.add_paragraph()
    name_title_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    name_title_para.paragraph_format.line_spacing = 1.15

    if title.strip():
        # If a title is provided, align name left and title right using a tab stop.
        name_title_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        tab_stops_name_title = name_title_para.paragraph_format.tab_stops
        tab_stops_name_title.add_tab_stop(_calculate_margin_end(doc), WD_TAB_ALIGNMENT.RIGHT)

        name_run = name_title_para.add_run(name)
        name_run.bold = True
        name_run.font.size = Pt(25)
        name_run.font.name = 'Cambria'
        
        title_run = name_title_para.add_run(f"\t{title}")
        title_run.bold = False
        title_run.font.size = Pt(18)
        title_run.font.name = 'Cambria'
    else:
        # If no title, center the name.
        name_title_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = name_title_para.add_run(name)
        name_run.bold = True
        name_run.font.size = Pt(25)
        name_run.font.name = 'Cambria'

    # Paragraph for Contact Info (location, email, phone, website)
    contact_parts = []
    if location.strip(): contact_parts.append(location.strip())
    if email.strip(): contact_parts.append(email.strip())
    if phone.strip(): contact_parts.append(phone.strip())
    
    # Add the website to contact_parts only if it's a valid URL, to ensure it's processed as a link
    if website.strip(): contact_parts.append(website.strip()) 

    if contact_parts:
        contact_para = doc.add_paragraph()
        contact_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        
        # Iterates through contact parts, making URLs clickable and adding separators.
        for i, part in enumerate(contact_parts):
            if part.startswith("http://") or part.startswith("https://"):
                # Clean the URL for display
                display_text_for_link = _clean_url_display(part)
                _add_hyperlink(doc, contact_para, part, display_text_for_link, size=10)
            else:
                run = contact_para.add_run(part)
                run.font.size = Pt(10)
                run.font.name = 'Cambria'
            
            # Adds a separator between contact details if it's not the last item.
            if i < len(contact_parts) - 1:
                run = contact_para.add_run(" | ")
                run.font.size = Pt(10)
                run.font.name = 'Cambria'
            
        contact_para.paragraph_format.space_after = Pt(0) # Keeps spacing tight in the header

    # Paragraph for LinkedIn and GitHub profiles.
    linkedin_github_para = doc.add_paragraph()
    linkedin_github_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    linkedin_github_para.paragraph_format.space_after = Pt(10) # Space after the entire header block
    linkedin_github_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    links_added = False
    if linkedin.strip():
        linkedin_display_text = _clean_url_display(linkedin.strip())
        _add_hyperlink(doc, linkedin_github_para, linkedin, linkedin_display_text, size=10)
        links_added = True
    
    if github.strip():
        if links_added:
            # Adds a separator if both LinkedIn and GitHub links are present.
            run = linkedin_github_para.add_run(" | ")
            run.font.size = Pt(10)
            run.font.name = 'Cambria'
        github_display_text = _clean_url_display(github.strip())
        _add_hyperlink(doc, linkedin_github_para, github, github_display_text, size=10)
        links_added = True


def add_summary_section(doc, summary_text):
    """Adds the Summary section of the resume."""
    if summary_text.strip():
        _add_section_title(doc, "Summary")
        p = doc.add_paragraph(summary_text)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(2)
        if p.runs: # Ensures font settings are applied if the paragraph has runs
            p.runs[0].font.size = Pt(10)
            p.runs[0].font.name = 'Cambria'

def add_education_section(doc, education_entries):
    """Adds the Education section, detailing university, degree, dates, GPA, and coursework."""
    # Filters for valid education entries to avoid adding empty sections.
    valid_education_entries = [
        edu for edu in education_entries
        if edu['university'].strip() or edu['degree'].strip() or
           edu['start_date'].strip() or edu['end_date'].strip() or
           edu['gpa'].strip() or edu['coursework'].strip()
    ]

    if valid_education_entries:
        _add_section_title(doc, "Education")
        for edu in valid_education_entries:
            # --- Line 1: University (Bold) <tabstop> Dates ---
            left_university_content = []
            if edu['university'].strip():
                # University text, bold
                left_university_content.append((edu['university'].strip(), True, False))
            
            # Add location to the university line if present
            if edu['location'].strip():
                if edu['university'].strip(): # If university is present, add comma
                    left_university_content.append((", ", False, False))
                left_university_content.append((edu['location'].strip(), False, False)) # Not bold, not italic

            right_dates_text_parts = []
            if edu['start_date'].strip(): 
                right_dates_text_parts.append(edu['start_date'].strip())
            if edu['end_date'].strip(): 
                right_dates_text_parts.append(edu['end_date'].strip())
            right_dates_text = " - ".join(right_dates_text_parts)

            # Use _add_left_right_paragraph for the first line (University/Location and Dates)
            if left_university_content or right_dates_text:
                _add_left_right_paragraph(doc, left_university_content, right_dates_text, None, left_size=10, right_size=10)


            # --- Line 2: Degree <tabstop> GPA (Italic) ---
            left_degree_content = []
            if edu['degree'].strip():
                left_degree_content.append((edu['degree'].strip(), False, False)) # Degree is not bold, not italic

            right_gpa_text = ""
            if edu['gpa'].strip():
                right_gpa_text = f"GPA: {edu['gpa'].strip()}"
            
            # Use _add_left_right_paragraph for the second line (Degree and GPA)
            # Pass right_italic=True to ensure GPA text is italicized.
            if left_degree_content or right_gpa_text:
                _add_left_right_paragraph(doc, left_degree_content, right_gpa_text, None, left_size=10, right_size=10, right_italic=True)


            # --- Line 3: Relevant Coursework ---
            coursework_text = edu['coursework'].strip()
            if coursework_text:
                coursework_para = doc.add_paragraph()
                coursework_para.paragraph_format.space_before = Pt(0)
                coursework_para.paragraph_format.space_after = Pt(2)
                coursework_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                
                # Clean up coursework text (it's a single string from text_area, but might have newlines)
                cleaned_coursework = ", ".join([line.strip() for line in coursework_text.splitlines() if line.strip()])
                
                coursework_run = coursework_para.add_run(f"Relevant Coursework: {cleaned_coursework}")
                coursework_run.font.size = Pt(10)
                coursework_run.font.name = 'Cambria'


def add_work_experience_section(doc, experiences):
    """Adds the Work Experience section, detailing job titles, companies, locations, dates, and responsibilities."""
    # Filters for valid experience entries.
    valid_experiences = [exp for exp in experiences if exp['job_title'].strip() and exp['company'].strip()]

    if valid_experiences:
        _add_section_title(doc, "Work Experience")
        for exp in valid_experiences:
            # Prepares left-aligned content (Job Title, Company, Location) with formatting.
            left_exp_content = []
            if exp['job_title'].strip():
                left_exp_content.append((exp['job_title'].strip(), True, False))
            
            if exp['company'].strip():
                if left_exp_content and not left_exp_content[-1][0].endswith(', '):
                    left_exp_content.append((", ", False, False))
                left_exp_content.append((exp['company'].strip(), False, False))
            
            if exp['location'].strip():
                if left_exp_content and not left_exp_content[-1][0].endswith(' - '):
                    left_exp_content.append((" - ", False, False))
                left_exp_content.append((exp['location'].strip(), False, True))

            # Prepares right-aligned dates content.
            right_exp_text_parts = []
            if exp['start_date'].strip(): right_exp_text_parts.append(exp['start_date'].strip())
            if exp['end_date'].strip(): right_exp_text_parts.append(exp['end_date'].strip())
            right_exp_text = " - ".join(right_exp_text_parts)

            _add_left_right_paragraph(doc, left_exp_content, right_exp_text, None, left_size=10, right_size=10)
            
            responsibilities_data = exp.get('responsibilities', [])
            
            # Ensure responsibilities_data is always a list of strings for iteration
            if isinstance(responsibilities_data, str):
                # If it's a single string, split it by lines and clean
                responsibilities_list = [line.strip() for line in responsibilities_data.splitlines() if line.strip()]
            elif isinstance(responsibilities_data, list):
                # If it's already a list, ensure elements are strings and not empty
                responsibilities_list = [str(item).strip() for item in responsibilities_data if str(item).strip()]
            else:
                responsibilities_list = [] # Default to empty list if unexpected type

            # Adds each responsibility as a bullet point.
            for desc in responsibilities_list:
                if desc: # Check again after stripping
                    if "O(n²)" in desc or "O(n log n)" in desc:
                        p = doc.add_paragraph(style='List Bullet')
                        p.paragraph_format.left_indent = Inches(0.25)
                        p.paragraph_format.first_line_indent = Inches(-0.25)
                        p.paragraph_format.space_after = Pt(2)
                        p.paragraph_format.line_before = Pt(2) # Corrected this line
                        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

                        # Splits the description and adds runs for normal text and special notations.
                        parts_n2 = desc.split("O(n²)")
                        for i, part_n2 in enumerate(parts_n2):
                            if part_n2:
                                parts_nlogn = part_n2.split("O(n log n)")
                                for j, part_nlogn in enumerate(parts_nlogn):
                                    if part_nlogn:
                                        run = p.add_run(part_nlogn)
                                    if j < len(parts_nlogn) - 1:
                                        run_nlogn = p.add_run('O(n log n)')
                            if i < len(parts_n2) - 1:
                                run_n2 = p.add_run('O(n²)')
                        
                        for run_item in p.runs:
                            run_item.font.name = 'Cambria'
                            run_item.font.size = Pt(10)
                    else:
                        _add_bullet_point(doc, desc)


def add_projects_section(doc, projects):
    """Adds the Projects section, detailing project titles, links, tech stack, and descriptions."""
    valid_projects = [proj for proj in projects if proj['title'].strip() or proj['description']] # Modified check for description

    if valid_projects:
        _add_section_title(doc, "Projects")
        for proj in valid_projects:
            # Adds the Project Title (left) and Live Demo Link (right).
            left_proj_title_content = [(proj['title'].strip(), True, False)]
            
            deployment_url = proj.get('deployment', '').strip()
            deployment_display_text = _clean_url_display(deployment_url) if deployment_url else ""
            
            _add_left_right_paragraph(doc, left_proj_title_content, deployment_display_text, deployment_url, left_size=10, right_size=10)
            
            # Adds Tech Stack (left) and GitHub Link (right) if either is present.
            tech_stack_text = proj.get('tech_stack', '').strip()
            github_url = proj.get('link', '').strip()
            github_display_text = _clean_url_display(github_url) if github_url else ""

            if tech_stack_text or github_url:
                left_tech_stack_content = []
                if tech_stack_text:
                    left_tech_stack_content.append((f"Tools Used: {tech_stack_text}", False, False))
                
                if left_tech_stack_content or github_url:
                    _add_left_right_paragraph(doc, left_tech_stack_content, github_display_text, github_url, left_size=10, right_size=10)

            project_description_data = proj.get('description') # Get the raw data

            if isinstance(project_description_data, str):
                description_list = [line.strip() for line in project_description_data.splitlines() if line.strip()]
            elif isinstance(project_description_data, list):
                description_list = [str(item).strip() for item in project_description_data if str(item).strip()]
            else:
                description_list = []

            for desc_line in description_list:
                if desc_line:
                    _add_bullet_point(doc, desc_line)

def add_skills_section(doc, technical_skills, soft_skills):
    """Adds the Skills section, categorizing technical and soft skills."""
    cleaned_tech_skills = [s.strip() for s in technical_skills if s.strip()]
    cleaned_soft_skills = [s.strip() for s in soft_skills if s.strip()]

    if cleaned_tech_skills or cleaned_soft_skills:
        _add_section_title(doc, "Skills")
        if cleaned_tech_skills:
            tech_para = doc.add_paragraph()
            tech_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            tech_para.paragraph_format.space_after = Pt(2)
            tech_para.paragraph_format.space_before = Pt(2)
            tech_para.add_run("Technical Skills: ").bold = True
            tech_para.runs[0].font.size = Pt(10)
            tech_para.runs[0].font.name = 'Cambria'
            
            if cleaned_tech_skills:
                tech_para.add_run("• " + " • ".join(cleaned_tech_skills))
            tech_para.runs[-1].font.size = Pt(10)
            tech_para.runs[-1].font.name = 'Cambria'

        if cleaned_soft_skills:
            soft_para = doc.add_paragraph()
            soft_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            soft_para.paragraph_format.space_after = Pt(2)
            soft_para.paragraph_format.space_before = Pt(2)
            soft_para.add_run("Soft Skills: ").bold = True
            soft_para.runs[0].font.size = Pt(10)
            soft_para.runs[0].font.name = 'Cambria'
            
            if cleaned_soft_skills:
                soft_para.add_run("• " + " • ".join(cleaned_soft_skills))
            soft_para.runs[-1].font.size = Pt(10)
            soft_para.runs[-1].font.name = 'Cambria'

def add_certifications_training_section(doc, certifications):
    """Adds the Certifications and Training section, including titles, issuers, and links."""
    valid_certifications = [cert for cert in certifications if cert['title'].strip()]

    if valid_certifications:
        _add_section_title(doc, "Certifications and Training")
        for cert in valid_certifications:
            left_cert_content = []
            if cert['title'].strip():
                left_cert_content.append((cert['title'].strip(), True, False))

            if cert.get('issuer', '').strip():
                if left_cert_content and not left_cert_content[-1][0].endswith(' '):
                    left_cert_content.append((" ", False, False))
                left_cert_content.append((f"({cert['issuer'].strip()})", False, False))

            cert_url = cert.get('link', '').strip()
            cert_display_text = _clean_url_display(cert_url) if cert_url else ""

            _add_left_right_paragraph(doc, left_cert_content, cert_display_text, cert_url, left_size=10, right_size=10)


def add_achievements_section(doc, achievements):
    """Adds the Achievements section as bullet points."""
    cleaned_achievements = [a.strip() for a in achievements if a.strip()]

    if cleaned_achievements:
        _add_section_title(doc, "Achievements")
        for achievement in cleaned_achievements:
            _add_bullet_point(doc, achievement)

def add_hobbies_section(doc, hobbies):
    """Adds the Hobbies section as a comma-separated list."""
    cleaned_hobbies = [h.strip() for h in hobbies if h.strip()]

    if cleaned_hobbies:
        _add_section_title(doc, "Hobbies")
        p = doc.add_paragraph(", ".join(cleaned_hobbies))
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(2)
        if p.runs:
            p.runs[0].font.size = Pt(10)
            p.runs[0].font.name = 'Cambria'

def generate_structured_resume(data, template_path=None):
    """
    Generates a complete resume document in DOCX format based on the provided structured data.
    """
    doc = Document()

    # Set narrow margins (0.5" all sides)
    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    # Force recalculation of cached margin end (important if this module is reused)
    global global_margin_end
    global_margin_end = None

    # Initializes document-wide default font and line spacing for consistency.
    _set_default_font_style(doc)
    _calculate_margin_end(doc) # Calculates the content area width for tab stops

    # Extracts data from the input dictionary for easier access.
    personal = data['personal']
    summary = data['summary']
    education = data['education']
    experience = data['experience']
    projects = data['projects']
    skills = data['skills']
    certifications = data['certifications']
    extras = data['achievements_hobbies']

    # Calls various section-specific functions to populate the document.
    add_header_section(doc,
        name=personal['name'],
        title=personal['title'],
        location=personal['location'],
        email=personal['email'],
        phone=personal['phone'],
        website=personal['website'],
        linkedin=personal['linkedin'],
        github=personal['github']
    )

    add_summary_section(doc, summary)
    add_education_section(doc, education)
    add_work_experience_section(doc, experience)
    add_projects_section(doc, projects)
    add_skills_section(doc,
        technical_skills=skills['technical'],
        soft_skills=skills['soft']
    )
    add_certifications_training_section(doc, certifications)
    add_achievements_section(doc, extras['achievements'])
    add_hobbies_section(doc, extras['hobbies'])

    # Saves the generated document to a BytesIO object and returns it.
    # This allows the document to be passed in memory (e.g., to Streamlit for download)
    # without needing to save it to a physical file first.
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio
